from docx import Document
from docx.oxml.ns import qn


def replace_in_paragraph(para, fields: dict):
    """
    Safely replace all {{placeholders}} in a paragraph.
    Handles cases where placeholders are split across multiple runs
    by merging all run text first, then redistributing.
    """
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    # Replace all placeholders in the merged text
    for key, value in fields.items():
        placeholder = f"{{{{{key}}}}}"
        replacement = str(value) if value else ""
        full_text = full_text.replace(placeholder, replacement)

    # Write merged text back into first run, clear the rest
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ""


def replace_all_placeholders(doc, fields: dict):
    """Replace placeholders in all paragraphs and table cells."""

    # Replace in top-level paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, fields)

    # Replace in tables (including nested tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, fields)


def find_paragraph_index(doc, search_text: str):
    """Return the index of the first paragraph containing search_text, or -1."""
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            return i
    return -1


def remove_paragraph(para):
    """Remove a paragraph element from the document."""
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def remove_special_conditions_section(doc):
    """
    Remove the entire Special Conditions section from the document.
    This removes the heading AND the placeholder/content paragraph underneath it.
    Only runs when special_conditions is null/empty.
    """
    paragraphs = doc.paragraphs
    to_remove = []
    inside_section = False

    for para in paragraphs:
        text = para.text.strip()
        text_lower = text.lower()

        # Detect the start of the special conditions section
        if "special conditions" in text_lower and not inside_section:
            inside_section = True
            to_remove.append(para)
            continue

        # Once inside, keep removing until we hit the placeholder or an empty line
        if inside_section:
            to_remove.append(para)
            if "{{special_conditions}}" in text or text == "":
                break

    for para in to_remove:
        remove_paragraph(para)


def generate_welcome_pack(fields: dict, template_path: str, output_path: str):
    """
    Load the Welcome Pack template, fill in all placeholders,
    optionally remove Special Conditions section, and save.
    """
    doc = Document(template_path)

    # Step 1: Replace all placeholders
    replace_all_placeholders(doc, fields)

    # Step 2: Remove special conditions section entirely if not present
    special = fields.get("special_conditions")
    if not special or str(special).strip().lower() in ("null", "none", "nil", "n/a", ""):
        remove_special_conditions_section(doc)

    # Step 3: Save output
    doc.save(output_path)
    print(f"Welcome pack saved to: {output_path}")