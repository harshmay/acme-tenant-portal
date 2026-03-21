from google import genai
import pdfplumber
import json
import os
from dotenv import load_dotenv
from docx import Document as DocxDocument


load_dotenv()
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))


FIELDS_PROMPT = """
You are extracting data from a lease agreement. Return ONLY a valid JSON object with exactly these fields:

{
  "tenant_name": "Full name(s). If joint tenancy, separate with ' and '",
  "property_address": "Full address including suburb, state, postcode",
  "lease_start_date": "DD/MM/YYYY",
  "lease_end_date": "DD/MM/YYYY",
  "rent_amount": "e.g. $2,400 per month or $1,200 per fortnight",
  "bond_amount": "Dollar amount e.g. $4,800",
  "num_occupants": "Number as stated in the lease",
  "pet_permission": "Not permitted OR description with any conditions",
  "special_conditions": "Full text if present, or null if none exist. If the lease says Nil, None, or No special conditions, return null.",
    "parking_included": "Return 'Not included' if no parking, otherwise describe the parking details.",
  "landlord_name": "As listed in parties section",
  "property_manager_name": "Contact person name",
  "property_manager_email": "Email address",
  "property_manager_phone": "Phone number"
}

Return ONLY the JSON. No explanation, no markdown, no code fences.

Lease text:
"""
def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    text = ""
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    # Extract tables — this is critical for structured leases
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    return text

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_fields(file_path: str) -> dict:
    if file_path.endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        text = extract_text_from_pdf(file_path)
    response = client.models.generate_content(
        model="gemini-2.5-flash-lite",
        contents= FIELDS_PROMPT + text
    )
    raw = response.text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())